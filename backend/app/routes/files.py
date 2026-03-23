"""
File Upload & Extraction API Routes
===================================
Handles document uploads for chat context enrichment.

Supported formats (phase 1):
- .txt / .md
- .pdf
- .docx
"""

from __future__ import annotations

import io
import logging
import os
import re
import tempfile
import time
import uuid
from pathlib import Path
from typing import Any, Callable, List, Optional

from fastapi import APIRouter, Depends, File, HTTPException, UploadFile, status
from pydantic import BaseModel

from ..models import User
from .auth_utils import get_current_user

router = APIRouter(prefix="/files", tags=["Files"])
logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB
MAX_FILES_PER_UPLOAD = 5
MAX_EXTRACTED_CHARS = 40_000
MAX_CHUNKS_PER_FILE = 24
TARGET_CHUNK_CHARS = 1_600
MAX_CHUNK_CHARS = 2_200
CHUNK_OVERLAP_CHARS = 180
MIN_PDF_TEXT_CHARS_FOR_OCR = 120
MAX_PDF_PAGES_FOR_OCR = 12

TEXT_MIME_TYPES = {"text/plain", "text/markdown"}
PDF_MIME_TYPES = {"application/pdf"}
DOCX_MIME_TYPES = {"application/vnd.openxmlformats-officedocument.wordprocessingml.document"}
SUPPORTED_EXTENSIONS = {".txt", ".md", ".markdown", ".pdf", ".docx"}

_HEADING_LINE_RE = re.compile(r"^(#{1,6}\s+.+|[A-Z0-9][A-Z0-9 .,:()/_-]{2,80}|.+:)$")
_OCR_UNAVAILABLE_RE = re.compile(r"(ocr not enabled|ocr unavailable|feature is off|not compiled with ocr)", re.IGNORECASE)


class FileChunkResponse(BaseModel):
    index: int
    label: Optional[str] = None
    text: str
    char_count: int
    truncated: bool = False


class FileUploadResponse(BaseModel):
    id: str
    filename: str
    mime_type: str
    size: int
    extracted_text: str
    text_chars: int
    truncated: bool
    chunks: List[FileChunkResponse]


class MultiFileUploadResponse(BaseModel):
    files: List[FileUploadResponse]
    count: int


def _normalize_text(text: str) -> str:
    normalized = (text or "").replace("\r\n", "\n").replace("\r", "\n").replace("\x00", "")
    normalized = "\n".join(line.rstrip() for line in normalized.split("\n"))
    normalized = re.sub(r"\n{3,}", "\n\n", normalized)
    return normalized.strip()


def _truncate_text(text: str, max_chars: int) -> tuple[str, bool]:
    if len(text) <= max_chars:
        return text, False
    return text[:max_chars].rstrip(), True


def _looks_like_heading(block: str) -> bool:
    text = (block or "").strip()
    if not text:
        return False
    if "\n" in text:
        return False
    if len(text) > 90:
        return False
    return bool(_HEADING_LINE_RE.match(text))


def _chunk_document_text(text: str) -> List[FileChunkResponse]:
    normalized = _normalize_text(text)
    if not normalized:
        return []

    raw_blocks = [part.strip() for part in normalized.split("\n\n") if part.strip()]
    blocks: list[str] = []
    for part in raw_blocks:
        lines = [line.rstrip() for line in part.splitlines()]
        if lines and _looks_like_heading(lines[0]):
            heading = lines[0].strip()
            body = "\n".join(line for line in lines[1:] if line.strip()).strip()
            blocks.append(heading)
            if body:
                blocks.append(body)
            continue
        blocks.append(part)
    if not blocks:
        blocks = [normalized]

    chunks: list[FileChunkResponse] = []
    current_parts: list[str] = []
    current_label: Optional[str] = None

    def flush_chunk() -> None:
        nonlocal current_parts, current_label
        if not current_parts or len(chunks) >= MAX_CHUNKS_PER_FILE:
            current_parts = []
            current_label = None
            return

        combined = "\n\n".join(part for part in current_parts if part).strip()
        if not combined:
            current_parts = []
            current_label = None
            return

        clipped, truncated = _truncate_text(combined, MAX_CHUNK_CHARS)
        chunk_index = len(chunks) + 1
        chunks.append(
            FileChunkResponse(
                index=chunk_index,
                label=current_label or f"Chunk {chunk_index}",
                text=clipped,
                char_count=len(combined),
                truncated=truncated,
            )
        )
        if len(chunks) >= MAX_CHUNKS_PER_FILE:
            current_parts = []
            current_label = None
            return

        overlap_source = clipped[-CHUNK_OVERLAP_CHARS:].strip()
        current_parts = [overlap_source] if overlap_source else []
        current_label = None

    for block in blocks:
        if _looks_like_heading(block):
            if current_parts:
                flush_chunk()
            current_label = block.lstrip("#").strip()[:80]
            current_parts.append(block)
            continue

        candidate_parts = [*current_parts, block]
        candidate_text = "\n\n".join(part for part in candidate_parts if part).strip()
        if candidate_text and len(candidate_text) > TARGET_CHUNK_CHARS and current_parts:
            flush_chunk()
            if _looks_like_heading(block):
                current_label = block.lstrip("#").strip()[:80]
            current_parts.append(block)
            continue

        current_parts.append(block)

    flush_chunk()

    if not chunks:
        clipped, truncated = _truncate_text(normalized, MAX_CHUNK_CHARS)
        return [
            FileChunkResponse(
                index=1,
                label="Chunk 1",
                text=clipped,
                char_count=len(normalized),
                truncated=truncated,
            )
        ]

    return chunks[:MAX_CHUNKS_PER_FILE]


def _extract_text_from_plain(data: bytes) -> str:
    return data.decode("utf-8", errors="replace")


def _pdf_texts_look_equivalent(primary: str, secondary: str) -> bool:
    first = re.sub(r"\s+", " ", _normalize_text(primary)).strip()
    second = re.sub(r"\s+", " ", _normalize_text(secondary)).strip()
    if not first or not second:
        return False
    shorter, longer = (first, second) if len(first) <= len(second) else (second, first)
    if shorter not in longer:
        return False
    return (len(shorter) / max(1, len(longer))) >= 0.75


def _stringify_pdf_table_cell(cell: Any) -> str:
    if cell is None:
        return ""
    if isinstance(cell, str):
        return _normalize_text(cell)
    if isinstance(cell, (int, float)):
        return str(cell)
    if isinstance(cell, dict):
        for key in ("text", "value", "content", "label"):
            value = cell.get(key)
            if isinstance(value, str) and value.strip():
                return _normalize_text(value)
        for key in ("cells", "values", "columns", "items"):
            value = cell.get(key)
            if isinstance(value, (list, tuple)):
                parts = [_stringify_pdf_table_cell(item) for item in value]
                parts = [part for part in parts if part]
                if parts:
                    return " ".join(parts)
        return ""

    for attr in ("text", "value", "content", "label"):
        value = getattr(cell, attr, None)
        if isinstance(value, str) and value.strip():
            return _normalize_text(value)

    if isinstance(cell, (list, tuple)):
        parts = [_stringify_pdf_table_cell(item) for item in cell]
        parts = [part for part in parts if part]
        return " ".join(parts)

    rendered = str(cell).strip()
    if rendered and not rendered.startswith("<"):
        return _normalize_text(rendered)
    return ""


def _coerce_pdf_table_rows(table: Any) -> list[list[str]]:
    if table is None:
        return []

    row_source: Any = None
    if isinstance(table, dict):
        for key in ("rows", "data", "values", "cells"):
            value = table.get(key)
            if isinstance(value, (list, tuple)):
                row_source = value
                break
    elif hasattr(table, "rows"):
        row_source = getattr(table, "rows", None)
    elif isinstance(table, (list, tuple)):
        row_source = table

    if not isinstance(row_source, (list, tuple)):
        return []

    rows: list[list[str]] = []
    for row in row_source:
        cell_source: Any
        if isinstance(row, dict):
            for key in ("cells", "values", "columns", "items"):
                value = row.get(key)
                if isinstance(value, (list, tuple)):
                    cell_source = value
                    break
            else:
                cell_source = list(row.values())
        elif hasattr(row, "cells"):
            cell_source = getattr(row, "cells", None) or []
        elif isinstance(row, (list, tuple)):
            cell_source = row
        else:
            cell_source = [row]

        if not isinstance(cell_source, (list, tuple)):
            cell_source = [cell_source]

        cells = [_stringify_pdf_table_cell(cell) for cell in cell_source]
        if any(cells):
            rows.append(cells)

    return rows


def _format_pdf_table_rows(rows: list[list[str]]) -> str:
    normalized_rows = [row for row in rows if any(cell.strip() for cell in row)]
    if not normalized_rows:
        return ""

    column_count = max(len(row) for row in normalized_rows)
    padded_rows = [row + [""] * (column_count - len(row)) for row in normalized_rows]
    header = padded_rows[0]
    body = padded_rows[1:]

    lines = [
        f"| {' | '.join(header)} |",
        f"| {' | '.join('---' for _ in range(column_count))} |",
    ]
    for row in body:
        lines.append(f"| {' | '.join(row)} |")
    return "\n".join(lines)


def _extract_pdf_table_text(doc: Any, page_index: int) -> str:
    try:
        raw_tables = doc.extract_tables(page_index)
    except Exception as exc:
        logger.debug("PDF table extraction skipped page=%d error=%s", page_index + 1, exc)
        return ""

    if not isinstance(raw_tables, (list, tuple)) or not raw_tables:
        return ""

    blocks: list[str] = []
    for table_index, table in enumerate(raw_tables, start=1):
        rows = _coerce_pdf_table_rows(table)
        if not rows:
            continue
        formatted = _format_pdf_table_rows(rows)
        if formatted:
            blocks.append(f"### Table {table_index}\n{formatted}")

    return "\n\n".join(blocks).strip()


def _extract_pdf_page_ocr(doc: Any, page_index: int) -> tuple[str, bool]:
    try:
        ocr_text = doc.extract_text_ocr(page_index)
    except Exception as exc:
        message = str(exc or "").strip()
        if _OCR_UNAVAILABLE_RE.search(message):
            logger.debug("PDF OCR unavailable page=%d error=%s", page_index + 1, message)
            return "", False
        logger.warning("PDF OCR failed page=%d error=%s", page_index + 1, message or exc)
        return "", True

    return _normalize_text(str(ocr_text or "")), True


def _extract_text_from_pdf_document(doc: Any) -> str:
    parts: list[str] = []
    ocr_supported = True

    for page_index in range(max(0, int(doc.page_count()))):
        page_text = _normalize_text(doc.extract_text(page_index) or "")
        page_sections: list[str] = []
        if page_text:
            page_sections.append(page_text)

        table_text = _extract_pdf_table_text(doc, page_index)
        if table_text:
            page_sections.append(table_text)

        should_try_ocr = (
            ocr_supported
            and page_index < MAX_PDF_PAGES_FOR_OCR
            and len(page_text) < MIN_PDF_TEXT_CHARS_FOR_OCR
        )
        if should_try_ocr:
            ocr_text, ocr_supported = _extract_pdf_page_ocr(doc, page_index)
            if ocr_text and not _pdf_texts_look_equivalent(page_text, ocr_text):
                if page_text:
                    page_sections.append(f"### OCR recovery\n{ocr_text}")
                else:
                    page_sections.append(ocr_text)

        page_body = "\n\n".join(section for section in page_sections if section).strip()
        if page_body:
            parts.append(f"## Page {page_index + 1}\n{page_body}")

    return "\n\n".join(parts)


def _extract_text_from_pdf_with_pdf_oxide(data: bytes) -> str:
    from pdf_oxide import PdfDocument

    temp_path: Optional[str] = None
    try:
        with tempfile.NamedTemporaryFile(mode="wb", suffix=".pdf", delete=False) as temp_pdf:
            temp_pdf.write(data)
            temp_path = temp_pdf.name

        doc = PdfDocument(temp_path)
        return _extract_text_from_pdf_document(doc)
    finally:
        if temp_path:
            try:
                os.remove(temp_path)
            except OSError:
                logger.warning("Failed to remove temporary PDF file: %s", temp_path)


def _extract_text_from_pdf(data: bytes) -> str:
    size = len(data)
    started_at = time.perf_counter()
    try:
        extracted = _extract_text_from_pdf_with_pdf_oxide(data)
        logger.info(
            "PDF extracted parser=%s size=%d chars=%d duration_ms=%.2f",
            "pdf_oxide",
            size,
            len(extracted),
            (time.perf_counter() - started_at) * 1000,
        )
        return extracted
    except Exception as exc:
        logger.warning(
            "PDF extraction failed parser=%s size=%d duration_ms=%.2f error=%s; falling back to pypdf",
            "pdf_oxide",
            size,
            (time.perf_counter() - started_at) * 1000,
            exc,
        )

    started_at = time.perf_counter()
    try:
        from pypdf import PdfReader
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="PDF parser is not available on server.",
        ) from exc

    reader = PdfReader(io.BytesIO(data))
    parts: list[str] = []
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = (page.extract_text() or "").strip()
        if page_text:
            parts.append(f"## Page {page_index}\n{page_text}")
    extracted = "\n\n".join(parts)
    logger.info(
        "PDF extracted parser=%s size=%d chars=%d duration_ms=%.2f",
        "pypdf",
        size,
        len(extracted),
        (time.perf_counter() - started_at) * 1000,
    )
    return extracted


def _extract_text_from_docx(data: bytes) -> str:
    try:
        from docx import Document
    except Exception as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="DOCX parser is not available on server.",
        ) from exc

    doc = Document(io.BytesIO(data))
    parts: list[str] = []

    for paragraph in doc.paragraphs:
        paragraph_text = (paragraph.text or "").strip()
        if not paragraph_text:
            continue

        style_name = (getattr(getattr(paragraph, "style", None), "name", "") or "").strip().lower()
        if style_name.startswith("heading"):
            match = re.search(r"(\d+)", style_name)
            level = max(1, min(6, int(match.group(1)) if match else 1))
            parts.append(f"{'#' * level} {paragraph_text}")
        else:
            parts.append(paragraph_text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _resolve_extractor(content_type: str, filename: str) -> Optional[Callable[[bytes], str]]:
    normalized_mime = (content_type or "").split(";", 1)[0].strip().lower()
    extension = Path(filename or "").suffix.lower()

    if normalized_mime in TEXT_MIME_TYPES or extension in {".txt", ".md", ".markdown"}:
        return _extract_text_from_plain
    if normalized_mime in PDF_MIME_TYPES or extension == ".pdf":
        return _extract_text_from_pdf
    if normalized_mime in DOCX_MIME_TYPES or extension == ".docx":
        return _extract_text_from_docx
    return None


@router.post("/upload", response_model=MultiFileUploadResponse)
async def upload_files(
    files: List[UploadFile] = File(...),
    current_user: User = Depends(get_current_user),
):
    """Upload and extract text from document files for chat context."""
    if len(files) > MAX_FILES_PER_UPLOAD:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"Maximum {MAX_FILES_PER_UPLOAD} files per upload.",
        )

    results: list[FileUploadResponse] = []

    for upload_file in files:
        filename = (upload_file.filename or "").strip() or f"file_{uuid.uuid4().hex[:8]}"
        extension = Path(filename).suffix.lower()
        content_type = (upload_file.content_type or "").split(";", 1)[0].strip().lower()

        extractor = _resolve_extractor(content_type, filename)
        if extractor is None or (extension and extension not in SUPPORTED_EXTENSIONS):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Unsupported file type: {filename}",
            )

        data = await upload_file.read()
        size = len(data)

        if size == 0:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Empty file: {filename}",
            )
        if size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File too large: {filename}. Maximum {MAX_FILE_SIZE} bytes.",
            )

        try:
            raw_text = extractor(data)
        except HTTPException:
            raise
        except Exception as exc:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"Failed to extract text from file: {filename}",
            ) from exc
        normalized_text = _normalize_text(raw_text)
        extracted_text, truncated = _truncate_text(normalized_text, MAX_EXTRACTED_CHARS)
        text_chars = len(normalized_text)
        chunks = _chunk_document_text(normalized_text)

        if not extracted_text:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"No extractable text found in file: {filename}",
            )

        results.append(
            FileUploadResponse(
                id=f"file_{uuid.uuid4().hex[:12]}",
                filename=os.path.basename(filename),
                mime_type=content_type or "application/octet-stream",
                size=size,
                extracted_text=extracted_text,
                text_chars=text_chars,
                truncated=truncated,
                chunks=chunks,
            )
        )

        logger.info(
            "File extracted user_id=%s filename=%s mime=%s size=%d chars=%d truncated=%s",
            current_user.id,
            filename,
            content_type or "application/octet-stream",
            size,
            text_chars,
            truncated,
        )

    return MultiFileUploadResponse(files=results, count=len(results))
