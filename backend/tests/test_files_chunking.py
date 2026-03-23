import io
import unittest

from docx import Document

from app.routes.files import (
    _chunk_document_text,
    _extract_text_from_docx,
    _extract_text_from_pdf_document,
)


class _FakePdfDocument:
    def __init__(self, pages):
        self._pages = pages

    def page_count(self):
        return len(self._pages)

    def extract_text(self, page_index):
        return self._pages[page_index].get("text", "")

    def extract_tables(self, page_index):
        return self._pages[page_index].get("tables", [])

    def extract_text_ocr(self, page_index):
        value = self._pages[page_index].get("ocr", "")
        if isinstance(value, Exception):
            raise value
        return value


class FilesChunkingTests(unittest.TestCase):
    def test_chunk_document_text_uses_page_markers_as_labels(self) -> None:
        text = (
            "## Page 1\nOverview and intro.\n\n"
            "## Page 2\nBilling details and repayment terms."
        )

        chunks = _chunk_document_text(text)

        self.assertEqual(len(chunks), 2)
        self.assertEqual(chunks[0].label, "Page 1")
        self.assertIn("Overview and intro.", chunks[0].text)
        self.assertEqual(chunks[1].label, "Page 2")
        self.assertIn("Billing details and repayment terms.", chunks[1].text)

    def test_extract_text_from_docx_preserves_heading_structure(self) -> None:
        buffer = io.BytesIO()
        document = Document()
        document.add_heading("Billing policy", level=1)
        document.add_paragraph("Installment payments are available for 12 months.")
        document.save(buffer)

        extracted = _extract_text_from_docx(buffer.getvalue())

        self.assertIn("# Billing policy", extracted)
        self.assertIn("Installment payments are available for 12 months.", extracted)

    def test_extract_text_from_pdf_document_formats_tables_and_ocr_recovery(self) -> None:
        doc = _FakePdfDocument(
            [
                {
                    "text": "Invoice",
                    "tables": [
                        {
                            "rows": [
                                ["Item", "Amount"],
                                ["Hosting", "$25"],
                                ["Domain", "$12"],
                            ]
                        }
                    ],
                    "ocr": "Invoice total 37 dollars paid on 2026-03-23.",
                }
            ]
        )

        extracted = _extract_text_from_pdf_document(doc)

        self.assertIn("## Page 1", extracted)
        self.assertIn("| Item | Amount |", extracted)
        self.assertIn("| Hosting | $25 |", extracted)
        self.assertIn("### OCR recovery", extracted)
        self.assertIn("Invoice total 37 dollars", extracted)

    def test_extract_text_from_pdf_document_skips_ocr_when_feature_is_unavailable(self) -> None:
        doc = _FakePdfDocument(
            [
                {
                    "text": "",
                    "tables": [],
                    "ocr": RuntimeError("OCR not enabled."),
                }
            ]
        )

        extracted = _extract_text_from_pdf_document(doc)

        self.assertEqual(extracted, "")


if __name__ == "__main__":
    unittest.main()
